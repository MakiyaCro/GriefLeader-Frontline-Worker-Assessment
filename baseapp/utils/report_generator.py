from docx import Document
from docx2pdf import convert
import os
from datetime import datetime
from django.conf import settings
from django.db.models import Avg
from baseapp.models import QuestionResponse, Attribute, AssessmentResponse

def get_benchmark_score_for_attribute(business, attribute):
    """Calculate benchmark score for a specific attribute"""
    benchmark_responses = AssessmentResponse.objects.filter(
        assessment__business=business,
        assessment__assessment_type='benchmark',
        assessment__completed=True
    )
    
    if not benchmark_responses.exists():
        return None
        
    total_score = 0
    responses = 0
    
    for response in benchmark_responses:
        score = response.get_score_for_attribute(attribute)
        if score is not None:
            total_score += score
            responses += 1
    
    return (total_score / responses) if responses > 0 else None

def generate_assessment_report(assessment_response):
    """
    Generate assessment report using Word template and convert to PDF.
    Handles placeholder replacement and score population in table.
    """
    # Get template path from utility directory
    current_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(current_dir, 'report_template.docx')
    
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Report template not found at {template_path}")
    
    # Define paths for temporary and final files
    reports_dir = os.path.join(settings.MEDIA_ROOT, 'assessment_reports')
    os.makedirs(reports_dir, exist_ok=True)
    
    temp_docx = os.path.join(reports_dir, f'temp_report_{assessment_response.assessment.unique_link}.docx')
    temp_pdf = os.path.join(reports_dir, f'temp_assessment_report_{assessment_response.assessment.unique_link}.pdf')
    final_pdf = os.path.join(reports_dir, f'assessment_report_{assessment_response.assessment.unique_link}.pdf')
    
    # Load the template
    doc = Document(template_path)
    
    # Get assessment data
    assessment = assessment_response.assessment
    completion_time = assessment_response.submitted_at - assessment.created_at
    
    # Replace placeholders in paragraphs
    placeholders = {
        '[candidate_name]': assessment.candidate_name,
        '[candidate_position]': assessment.position,
        '[submitted_at]': assessment_response.submitted_at.strftime('%Y-%m-%d'),
        '[completion_time]': str(completion_time).split('.')[0],  # Remove microseconds
        '[manager_name]': assessment.manager_name,
        '[region]': assessment.region
    }
    
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
    
    # Handle table scores
    if len(doc.tables) > 0:
        table = doc.tables[0]
        
        # Get all active attributes
        attributes = Attribute.objects.filter(
            business=assessment.business,
            active=True
        ).order_by('order')
        
        # Start from row 1 to skip header
        for row in table.rows[1:]:
            # Get attribute name from first column
            attribute_name = row.cells[0].text.strip()
            # Remove any formatting markers
            attribute_name = ''.join(c for c in attribute_name if c not in ('*', '[', ']', '{', '}', '.', ' '))
            
            try:
                # Find the matching attribute (case-insensitive)
                attr = next(
                    (a for a in attributes if a.name.lower().replace(' ', '') == attribute_name.lower()),
                    None
                )
                
                if attr:
                    # Calculate and format candidate score
                    candidate_score = assessment_response.get_score_for_attribute(attr)
                    
                    # Get benchmark score
                    benchmark_score = get_benchmark_score_for_attribute(assessment.business, attr)
                    
                    # Update the "Candidate" column (index 2)
                    if len(row.cells) > 2:
                        score_cell = row.cells[2]
                        score_cell.text = f"{candidate_score:.1f}%" if candidate_score is not None else "N/A"
                    
                    # Update the "Benchmark" column (index 3)
                    if len(row.cells) > 3:
                        benchmark_cell = row.cells[3]
                        benchmark_cell.text = f"{benchmark_score:.1f}%" if benchmark_score is not None else "N/A"
                        
            except Exception as e:
                print(f"Error processing attribute {attribute_name}: {str(e)}")
                continue
    
    # Save and convert the document
    try:
        # Save the Word document
        doc.save(temp_docx)
        
        # Convert to PDF
        convert(temp_docx, temp_pdf)
        
        # Move temporary PDF to final location
        if os.path.exists(temp_pdf):
            if os.path.exists(final_pdf):
                os.remove(final_pdf)
            os.rename(temp_pdf, final_pdf)
            
        return final_pdf
        
    finally:
        # Clean up temporary files
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
        if os.path.exists(temp_pdf):
            os.remove(temp_pdf)